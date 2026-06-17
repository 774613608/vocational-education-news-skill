#!/usr/bin/env python3
"""速览模式 Word — 全源覆盖 + 严格日期过滤 06.10-06.16"""
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml.shared import OxmlElement
from docx.opc.constants import RELATIONSHIP_TYPE as RT
from datetime import datetime

date_range = "2026年6月10日 — 6月16日（严格过滤）"
generated = datetime.now().strftime("%Y-%m-%d %H:%M")

SOURCES = [
    {
        "name": "教育部 (moe.gov.cn)",
        "color": "1a3c6d",
        "items": [
            {"id":1,"title":"国务院审议通过《教育发展十五五规划》","date":"06.11","oneliner":"李强主持，教育优先发展、健全学龄人口适配机制","url":"http://hudong.moe.gov.cn/jyb_xwfb/s6052/moe_838/202606/t20260612_1440488.html"},
            {"id":2,"title":"习近平《一体推进教育科技人才发展》发表","date":"06.15","oneliner":"统筹职教高教继教，职普融通产教融合科教融汇","url":"https://hudong.moe.gov.cn/jyb_xwfb/s6052/moe_838/202606/t20260615_1440767.html"},
            {"id":3,"title":"怀进鹏会见葡萄牙教育部长","date":"06.11","oneliner":"推动海洋、生命健康、数字技术、AI等领域教育合作","url":"https://hudong.moe.gov.cn/jyb_zzjg/huodong/202606/t20260612_1440569.html"},
            {"id":4,"title":"教育舆论宣传工作会议在雄安召开","date":"06.11-12","oneliner":"王光彦出席，推动AI与教育宣传工作深度融合","url":"https://hudong.moe.gov.cn/jyb_zzjg/huodong/202606/t20260612_1440565.html"},
        ]
    },
    {
        "name": "人民日报/人民网",
        "color": "d93025",
        "items": [
            {"id":5,"title":"习近平：一体推进教育科技人才发展（头版）","date":"06.16","oneliner":"教育科技人才三位一体，培养大国工匠能工巧匠","url":"https://paper.people.com.cn/rmrb/pc/content/202606/16/content_30163258.html"},
            {"id":6,"title":"丁仲礼率队在江苏检查职教法实施","date":"06.10-13","oneliner":"南京无锡调研，压实职普融通产教融合","url":"http://politics.people.com.cn/n1/2026/0613/c1001-40739500.html"},
            {"id":7,"title":"评论：职业教育如何向新而变","date":"06.12","oneliner":"烧烤/螺蛳粉/大疆学院三条路径对接产业新需求","url":"https://paper.people.com.cn/rmrb/pc/content/202606/12/content_30162655.html"},
        ]
    },
    {
        "name": "人社部 — 世赛与技能培训 (mohrss.gov.cn)",
        "color": "e37400",
        "items": [
            {"id":8,"title":"第48届世赛全力冲刺备战","date":"06.16","oneliner":"71选手64项目，新增无人机等7赛项，9月上海","url":"http://chinajob.mohrss.gov.cn/h5/c/2026-06-16/558115.shtml"},
            {"id":9,"title":"世赛倒计时100天系列活动举办","date":"06.14","oneliner":"1+10+100联动，李子柒任推广大使，技能博览会","url":"http://chinajob.mohrss.gov.cn/h5/c/2026-06-16/558117.shtml"},
            {"id":10,"title":"广州世赛基地探访：AI融入训练","date":"06.12","oneliner":"冠军带冠军+AI辅助赛题+云端测试平台","url":"http://chinajob.mohrss.gov.cn/h5/c/2026-06-12/556924.shtml"},
        ]
    },
    {
        "name": "现代高等职业技术教育网 (tech.net.cn)",
        "color": "2e7d32",
        "items": [
            {"id":11,"title":"湖南工程职院：无界教学湖南方案","date":"06.16","oneliner":"突破四大边界+BIM复刻2589节点，效率3倍提升","url":"https://www.tech.net.cn/news/show-108280.html"},
            {"id":12,"title":"哈尔滨科学技术职院：校社共建智慧康养","date":"06.15","oneliner":"AI赋能康养专业群，社区真实服务驱动学习","url":"https://www.tech.net.cn/news/show-108262.html"},
            {"id":13,"title":"江西制造职院：七能模式培育现场工程师","date":"06.15","oneliner":"精操作懂工艺能装调会管控可排故会改善能带队","url":"https://www.tech.net.cn/news/show-108263.html"},
            {"id":14,"title":"四川交通职院：培养多制式列车检修工匠","date":"06.15","oneliner":"多维发力，产教协同培养轨道交通高技能人才","url":"https://www.tech.net.cn/news/show-108269.html"},
            {"id":15,"title":"南京工业职大：探路产教融合深度嵌入","date":"06.15","oneliner":"共同体推动校企从浅层合作走向深度嵌入","url":"https://www.tech.net.cn/news/show-108267.html"},
            {"id":16,"title":"重庆电力高专：学中做做中学里练真功","date":"06.12","oneliner":"六维递进产教协同，VR仿真缩短理解时间40%","url":"https://www.tech.net.cn/news/show-108245.html"},
        ]
    },
    {
        "name": "中国教育在线 (eol.cn) — 本轮新增",
        "color": "1a73e8",
        "items": [
            {"id":17,"title":"明昱智培数字化技能实训就业平台启动","date":"06.15","oneliner":"AI+智能制造+工业机器人，培训实训考证就业全链条","url":"https://gaokao.eol.cn/gzxy/yxdt/202606/t20260615_2745157.shtml"},
            {"id":18,"title":"机械行业职业教育高质量发展座谈会","date":"06.13","oneliner":"聚焦职业本科建设、双高计划、产教融合、职教国际化","url":"https://chongqing.eol.cn/cqzy/202606/t20260613_2744805.shtml"},
            {"id":19,"title":"山东圣翰财贸职院校企合作双选会","date":"06.15","oneliner":"200余家企业5000+岗位，初步达成意向2327个","url":"https://shandong.eol.cn/sdzy/202606/t20260615_2744945.shtml"},
            {"id":20,"title":"广州华夏职院：五方联动激活乡村振兴","date":"06.13","oneliner":"校政行企社党建联建，产教融合+乡村振兴协同","url":"https://guangdong.eol.cn/gdgd/202606/t20260613_2744812.shtml"},
            {"id":21,"title":"全国汽车流通行业职业技能竞赛决赛","date":"06.13","oneliner":"50所院校95名选手在达州职院同台竞技","url":"https://sichuan.eol.cn/sczy/202606/t20260613_2744803.shtml"},
            {"id":22,"title":"遂宁职业学院首届毕业典礼暨招聘会","date":"06.16","oneliner":"1601名首届毕业生+百日冲刺专场招聘同步","url":"https://sichuan.eol.cn/sczy/202606/t20260616_2745403.shtml"},
        ]
    },
    {
        "name": "世界职业院校技能大赛 — 各省备赛动态（中文名搜索）",
        "color": "6a1b9a",
        "items": [
            {"id":23,"title":"商丘职院召开世校赛备赛座谈会","date":"06.15","oneliner":"副校长潘自舒部署暑期集训，精准备赛守住底线","url":"https://jiaowu.sqzy.edu.cn/2026/0616/c145a71640/page.htm"},
            {"id":24,"title":"滁州职院专家团队指导世校赛选拔赛","date":"06.14","oneliner":"针对实操细节技术规范流程把控全方位点评","url":"https://www.chzc.edu.cn/tmx/info/1022/36323.htm"},
            {"id":25,"title":"浙江金融职院召开世校赛备赛推进会","date":"06.12","oneliner":"校长提出技能体能心理三位一体六点要求","url":"https://www.zfc.edu.cn/2026/0612/c2771a66894/page.htm"},
            {"id":26,"title":"安徽工业经济职院世校赛动员会","date":"06.10","oneliner":"23支队伍入围省选拔赛，5支队直获世校赛资格","url":"https://www.uta.edu.cn/2026/0610/c1111a172111/page.htm"},
            {"id":27,"title":"开封大学召开世校赛备赛工作会","date":"06.10","oneliner":"副校长李治提五点要求：创新性熟练度心理测试","url":"http://www.kfu.edu.cn/news/info/1004/18257.htm"},
            {"id":28,"title":"重庆安全技术职院指导世校赛备赛","date":"06.10","oneliner":"副校长深入应急管理学院，对接行业技能标准","url":"http://www.cqvist.net/index.php?c=show&id=6183"},
        ]
    },
]

# ============================================================
doc = Document()

def set_run_font(run, ascii_font='Calibri', ea_font='微软雅黑', size=Pt(11), bold=False, color=None):
    run.font.size = size; run.bold = bold; run.font.name = ascii_font
    if color: run.font.color.rgb = color
    rPr = run._r.get_or_add_rPr()
    rFonts = rPr.find(qn('w:rFonts'))
    if rFonts is None:
        rFonts = OxmlElement('w:rFonts'); rPr.insert(0, rFonts)
    rFonts.set(qn('w:ascii'), ascii_font); rFonts.set(qn('w:hAnsi'), ascii_font)
    rFonts.set(qn('w:eastAsia'), ea_font)

def add_hyperlink(paragraph, url, text):
    part = paragraph.part; r_id = part.relate_to(url, RT.HYPERLINK, is_external=True)
    hyperlink = OxmlElement('w:hyperlink'); hyperlink.set(qn('r:id'), r_id)
    new_run = OxmlElement('w:r'); rPr = OxmlElement('w:rPr')
    for tag, val in [('w:color','0563C1'),('w:u','single'),('w:sz','16')]:
        el = OxmlElement(tag); el.set(qn('w:val'), val); rPr.append(el)
    new_run.append(rPr); new_run.text = text; hyperlink.append(new_run)
    paragraph._p.append(hyperlink)

def set_cell(cell, text, size=Pt(9), bold=False, color=None):
    for p in cell.paragraphs: p.clear()
    p = cell.paragraphs[0]; p.paragraph_format.space_before = Pt(1); p.paragraph_format.space_after = Pt(1)
    run = p.add_run(str(text)); set_run_font(run, size=size, bold=bold, color=color)

def link_cell(cell, url, text):
    for p in cell.paragraphs: p.clear()
    p = cell.paragraphs[0]; p.paragraph_format.space_before = Pt(1); p.paragraph_format.space_after = Pt(1)
    add_hyperlink(p, url, text)

# Title
title = doc.add_heading('最近一周教育新闻动态速览', level=0)
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
for r in title.runs: set_run_font(r, size=Pt(22), bold=True, color=RGBColor(0x1A,0x3C,0x6D))

sub = doc.add_paragraph(); sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = sub.add_run(f'数据范围：{date_range}    生成时间：{generated}    速览模式')
set_run_font(run, size=Pt(10), color=RGBColor(0x66,0x66,0x66))

total = sum(len(s['items']) for s in SOURCES)
sub2 = doc.add_paragraph(); sub2.alignment = WD_ALIGN_PARAGRAPH.CENTER
run2 = sub2.add_run(f'共 {total} 条 | {len(SOURCES)} 类信息源 | 覆盖教育部/人社部/人民日报/tech.net.cn/eol.cn/世校赛/civte + 重点院校 | 严格日期过滤')
set_run_font(run2, size=Pt(9), color=RGBColor(0x99,0x99,0x99))

sub3 = doc.add_paragraph(); sub3.alignment = WD_ALIGN_PARAGRAPH.CENTER
run3 = sub3.add_run('🟢 6源可检索(eol.cn/世校赛本轮新增) | 🟡 civte已爬取但本周无更新 | 🔴 cvae=Vue SPA无法爬取 | 链接均可点击')
set_run_font(run3, size=Pt(8), color=RGBColor(0xAA,0xAA,0xAA))
doc.add_paragraph('')

for src in SOURCES:
    ch = src['color']
    h = doc.add_heading(f"{src['name']} — {len(src['items'])}条", level=2)
    for r in h.runs: set_run_font(r, size=Pt(14), bold=True, color=RGBColor(int(ch[0:2],16),int(ch[2:4],16),int(ch[4:6],16)))

    table = doc.add_table(rows=len(src['items'])+1, cols=4, style='Table Grid'); table.autofit = True
    for i, hdr in enumerate(['#','标题','日期','一句话速览（点击跳转原文）']):
        c = table.rows[0].cells[i]; set_cell(c, hdr, size=Pt(9), bold=True, color=RGBColor(0xFF,0xFF,0xFF))
        sh = OxmlElement('w:shd'); sh.set(qn('w:fill'), ch); sh.set(qn('w:val'), 'clear')
        c._tc.get_or_add_tcPr().append(sh)

    for ri, item in enumerate(src['items']):
        row = table.rows[ri+1]
        set_cell(row.cells[0], str(item['id']), size=Pt(8), color=RGBColor(0x88,0x88,0x88))
        set_cell(row.cells[1], item['title'], size=Pt(9), bold=True)
        set_cell(row.cells[2], item['date'], size=Pt(8), color=RGBColor(0xE3,0x74,0x00))
        link_cell(row.cells[3], item['url'], item['oneliner'])
        if ri % 2 == 0:
            for cell in row.cells:
                sh = OxmlElement('w:shd'); sh.set(qn('w:fill'), 'F5F7FA'); sh.set(qn('w:val'), 'clear')
                cell._tc.get_or_add_tcPr().append(sh)
    for row in table.rows:
        row.cells[0].width = Cm(1.0); row.cells[1].width = Cm(6.0)
        row.cells[2].width = Cm(2.5); row.cells[3].width = Cm(7.5)
    doc.add_paragraph('')

doc.add_paragraph('')
ps = doc.add_paragraph(); ps.alignment = WD_ALIGN_PARAGRAPH.CENTER
rs = ps.add_run(f'📊 共 {total} 条（全部严格在06.10-06.16窗口内） | {len(SOURCES)} 类信息源 | 17类源覆盖14类(1不可爬+1本周无更新+2国际未涉及)')
set_run_font(rs, size=Pt(9), color=RGBColor(0x66,0x66,0x66))

pn = doc.add_paragraph(); pn.alignment = WD_ALIGN_PARAGRAPH.CENTER
rn = pn.add_run('注：civte.edu.cn 已通过curl直爬验证可访问，但工作动态最新为06.01，本周窗口内无更新。cvae.com.cn 为Vue SPA，curl无法渲染。')
set_run_font(rn, size=Pt(8), color=RGBColor(0x88,0x88,0x88))

fp = doc.add_paragraph(); fp.alignment = WD_ALIGN_PARAGRAPH.CENTER
rf = fp.add_run('🤖 速览模式 v3.0 | 全源覆盖 | 严格日期过滤 | 基于《职教AI知识库-信息源网址与栏目映射》| Generated by Claude Code')
set_run_font(rf, size=Pt(8), color=RGBColor(0xAA,0xAA,0xAA))

out = "examples/education_news_quickview.docx"
doc.save(out)
print(f"✅ {out}")
print(f"{len(SOURCES)} 源、{total} 条（全部严格06.10-06.16）")
