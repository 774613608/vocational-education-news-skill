#!/usr/bin/env python3
"""主题模式 Word — 全源覆盖 + 严格 06.10-06.16"""
import json
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml.shared import OxmlElement
from docx.opc.constants import RELATIONSHIP_TYPE as RT
from datetime import datetime

date_range = "2026年6月10日 — 6月16日"
generated = datetime.now().strftime("%Y-%m-%d %H:%M")

DATA = r'''
[
  {
    "theme": "一、顶层政策与战略部署",
    "color": "1a3c6d",
    "items": [
      {
        "title": "国务院常务会议审议通过《教育发展十五五规划》",
        "date": "06.11", "source": "新华社 / 教育部",
        "url": "http://hudong.moe.gov.cn/jyb_xwfb/s6052/moe_838/202606/t20260612_1440488.html",
        "summary": "李强总理主持国务院常务会议审议通过，强调教育优先发展、优化结构、提高质量，健全适应学龄人口变化的教育资源配置机制，强化教育对科技和人才的支撑作用，塑造立德树人新格局。"
      },
      {
        "title": "习近平《一体推进教育科技人才发展》重要文章发表",
        "date": "06.15-16", "source": "求是 / 人民日报",
        "url": "https://paper.people.com.cn/rmrb/pc/content/202606/16/content_30163258.html",
        "summary": "《求是》杂志第12期+人民日报头版发表，明确统筹职业教育、高等教育、继续教育，推进职普融通、产教融合、科教融汇，源源不断培养高素质技术技能人才、大国工匠、能工巧匠。"
      },
      {
        "title": "2026年教育舆论宣传工作会议在雄安召开",
        "date": "06.11-12", "source": "教育部",
        "url": "https://hudong.moe.gov.cn/jyb_zzjg/huodong/202606/t20260612_1440565.html",
        "summary": "教育部副部长王光彦出席并讲话，强调2026年是十五五开局之年，要强化数智创新，推动人工智能等新技术与教育舆论宣传工作深度融合。"
      }
    ]
  },
  {
    "theme": "二、职业教育法执法检查",
    "color": "d93025",
    "items": [
      {
        "title": "丁仲礼率执法检查组在江苏实地检查",
        "date": "06.10-13", "source": "人民日报",
        "url": "http://politics.people.com.cn/n1/2026/0613/c1001-40739500.html",
        "summary": "全国人大常委会副委员长丁仲礼赴南京、无锡实地调研职业院校和企业，强调建设职普融通、产教融合现代职业教育体系，压实各方法定职责，保障资源配置与政策供给。"
      },
      {
        "title": "人民日报评论：职业教育如何向新而变",
        "date": "06.12", "source": "人民日报",
        "url": "https://paper.people.com.cn/rmrb/pc/content/202606/12/content_30162655.html",
        "summary": "以岳阳烧烤学院、柳州螺蛳粉产业学院、深圳信息职业技术大学×大疆等为例，提出职业教育对接经济社会三条路径：特色文化赋能经济、对接新技术新业态、读懂百姓需求清单。教育部已建立职教新专业快速响应通道，重点布局低空经济、人工智能、高端装备。"
      }
    ]
  },
  {
    "theme": "三、产教融合与实训实践",
    "color": "2e7d32",
    "items": [
      {
        "title": "湖南工程职业技术学院：无界教学重构人才培养湖南方案",
        "date": "06.16", "source": "中国青年报 / tech.net.cn",
        "url": "https://www.tech.net.cn/news/show-108280.html",
        "summary": "提出无界教学理论框架，突破时间、认知、主体、空间四大边界，培养懂智能操作、能理解数据、会智能协同的复合型智建师。联合72家企业共建，BIM+VR复刻2589个施工节点，学生整体学习效率提升3倍。"
      },
      {
        "title": "南京工业职业技术大学：探路产教融合深度嵌入",
        "date": "06.15", "source": "中国教育报 / tech.net.cn",
        "url": "https://www.tech.net.cn/news/show-108267.html",
        "summary": "以共同体建设推动校企从浅层合作走向深度嵌入，与北京精雕等行业龙头共建产业学院，实现人才培养标准与产业需求精准对接。"
      },
      {
        "title": "重庆电力高等专科学校：在学中做、做中学里练真功",
        "date": "06.12", "source": "光明日报 / tech.net.cn",
        "url": "https://www.tech.net.cn/news/show-108245.html",
        "summary": "探索六维递进产教协同育人模式，利用VR虚拟仿真实训将电厂搬进教室，学生原理理解时间缩短40%，操作规范率提升35%。"
      },
      {
        "title": "明昱智培数字化技能实训就业平台正式启动",
        "date": "06.15", "source": "中国教育在线 (eol.cn)",
        "url": "https://gaokao.eol.cn/gzxy/yxdt/202606/t20260615_2745157.shtml",
        "summary": "在中科院人才交流开发中心启动，聚焦AI应用、智能制造、工业机器人等领域，打造培训-实训-考证-就业全链条服务体系。"
      },
      {
        "title": "机械行业职业教育高质量发展专题座谈会",
        "date": "06.13", "source": "中国教育在线 (eol.cn)",
        "url": "https://chongqing.eol.cn/cqzy/202606/t20260613_2744805.shtml",
        "summary": "在重庆召开，恰逢重庆工业职业技术大学建校70周年。会议聚焦职业本科建设、双高计划、产教融合、职教国际化等议题深入研讨。"
      }
    ]
  },
  {
    "theme": "四、课堂教学模式创新",
    "color": "1a73e8",
    "items": [
      {
        "title": "哈尔滨科学技术职业学院：校社共建培养智慧康养高技能人才",
        "date": "06.15", "source": "光明日报 / tech.net.cn",
        "url": "https://www.tech.net.cn/news/show-108262.html",
        "summary": "AI赋能智慧康养专业群建设，以社区真实服务需求驱动学习，学生在真实场景中掌握健康管理、康复护理等核心技能。"
      },
      {
        "title": "江西制造职业技术学院：七能模式培育现场工程师",
        "date": "06.15", "source": "光明日报 / tech.net.cn",
        "url": "https://www.tech.net.cn/news/show-108263.html",
        "summary": "创新七能培养模式——精操作、懂工艺、能装调、会管控、可排故、会改善、能带队，面向先进制造业培育现场工程师。"
      },
      {
        "title": "四川交通职业技术学院：多维发力培养多制式列车检修工匠",
        "date": "06.15", "source": "中国教育报 / tech.net.cn",
        "url": "https://www.tech.net.cn/news/show-108269.html",
        "summary": "聚焦轨道交通产业需求，构建多制式列车检修人才培养体系，产教协同培养覆盖地铁、高铁、城际等多种制式的检修工匠。"
      }
    ]
  },
  {
    "theme": "五、世界技能大赛与技能培训",
    "color": "e37400",
    "items": [
      {
        "title": "第48届世赛全力冲刺：71选手64项目覆盖",
        "date": "06.14-16", "source": "人社部",
        "url": "http://chinajob.mohrss.gov.cn/h5/c/2026-06-16/558115.shtml",
        "summary": "第48届世赛将于9月22-27日在上海举行。中国派出71名选手覆盖全部64个项目，新增无人机系统、数字交互媒体设计、软件测试、轨道车辆技术等7赛项。倒计时100天，上海1+10+100联动，李子柒出任推广大使。"
      },
      {
        "title": "广州世赛集训基地探访：AI融入训练体系",
        "date": "06.12", "source": "人社部",
        "url": "http://chinajob.mohrss.gov.cn/h5/c/2026-06-12/556924.shtml",
        "summary": "广州工贸技师学院移动应用开发与软件测试赛项，往届金牌得主转型教练实现冠军带冠军传承，AI辅助赛题设计、智能测试工具和云端测试平台投入使用。"
      },
      {
        "title": "世赛倒计时100天系列活动举办",
        "date": "06.14", "source": "人社部",
        "url": "http://chinajob.mohrss.gov.cn/h5/c/2026-06-16/558117.shtml",
        "summary": "上海采用1个主会场+10个分会场+百家技能单位联动模式，世界技能博览会将推出近百种技能体验项目。"
      }
    ]
  },
  {
    "theme": "六、世校赛省级选拔赛密集备战",
    "color": "6a1b9a",
    "items": [
      {
        "title": "安徽省：多校密集召开选拔赛备赛会",
        "date": "06.10-15", "source": "各院校官网",
        "url": "https://www.uta.edu.cn/2026/0610/c1111a172111/page.htm",
        "summary": "安徽工业经济职院23支队伍入围省选拔赛、5支直获世校赛资格。商丘职院副校长部署暑期集训计划。滁州职院专家团队针对实操细节、技术规范、流程把控全方位指导。浙江金融职院校领导提出技能体能心理三位一体六点要求。"
      },
      {
        "title": "河南/重庆/湖北：多省同步推进备赛",
        "date": "06.10-12", "source": "各院校官网",
        "url": "http://www.kfu.edu.cn/news/info/1004/18257.htm",
        "summary": "开封大学副校长提五点要求：提升作品创新性、强化训练熟练度、加强心理压力测试、确保方案全面性、深入研究比赛细则。重庆安全技术职院副校长深入应急管理学院实地指导，要求对接实际工作场景和行业技能操作标准。"
      }
    ]
  },
  {
    "theme": "七、院校就业与校企合作",
    "color": "1a3c6d",
    "items": [
      {
        "title": "山东圣翰财贸职业学院校企合作双选会",
        "date": "06.15", "source": "中国教育在线 (eol.cn)",
        "url": "https://shandong.eol.cn/sdzy/202606/t20260615_2744945.shtml",
        "summary": "200余家用人单位参会，提供5000余个岗位，初步达成就业意向2327个，政校企协同促就业成效显著。"
      },
      {
        "title": "广州华夏职业学院：五方联动激活乡村振兴",
        "date": "06.13", "source": "中国教育在线 (eol.cn)",
        "url": "https://guangdong.eol.cn/gdgd/202606/t20260613_2744812.shtml",
        "summary": "校政行企社党建联建，以红色引擎激活乡村振兴一池春水，推动产教融合与乡村振兴协同发展。"
      },
      {
        "title": "全国汽车流通行业职业技能竞赛决赛在达州举行",
        "date": "06.13", "source": "中国教育在线 (eol.cn)",
        "url": "https://sichuan.eol.cn/sczy/202606/t20260613_2744803.shtml",
        "summary": "第五届全国汽车流通行业职业技能竞赛决赛汽配销售经理人赛项，50所院校95名选手同台竞技。"
      },
      {
        "title": "遂宁职业学院首届毕业典礼暨百日冲刺招聘会",
        "date": "06.16", "source": "中国教育在线 (eol.cn)",
        "url": "https://sichuan.eol.cn/sczy/202606/t20260616_2745403.shtml",
        "summary": "1601名首届毕业生完成学业，典礼同步举办百日冲刺专场招聘会，助力毕业生顺利就业。"
      }
    ]
  },
  {
    "theme": "八、国际合作",
    "color": "d93025",
    "items": [
      {
        "title": "怀进鹏会见葡萄牙教育、科学和创新部部长",
        "date": "06.11", "source": "教育部",
        "url": "https://hudong.moe.gov.cn/jyb_zzjg/huodong/202606/t20260612_1440569.html",
        "summary": "教育部部长怀进鹏在京会见葡方部长费尔南多·亚历山大，就深化两国教育交流合作交换意见，推动海洋、生命健康、数字技术、人工智能等领域实质性合作。"
      }
    ]
  }
]
'''

sections = json.loads(DATA)

doc = Document()

def set_run_font(run, ascii_font='Calibri', ea_font='微软雅黑', size=Pt(11), bold=False, color=None):
    run.font.size = size; run.bold = bold; run.font.name = ascii_font
    if color: run.font.color.rgb = color
    rPr = run._r.get_or_add_rPr()
    rFonts = rPr.find(qn('w:rFonts'))
    if rFonts is None:
        rFonts = OxmlElement('w:rFonts'); rPr.insert(0, rFonts)
    rFonts.set(qn('w:ascii'), ascii_font); rFonts.set(qn('w:hAnsi'), ascii_font)
    rFonts.set(qn('w:eastAsia'), ea_font)

for name in ['Heading 1', 'Heading 2']:
    hs = doc.styles[name]; hs.font.name = 'Calibri'

title = doc.add_heading('最近一周教育新闻动态汇总（按主题分类）', level=0)
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
for r in title.runs: set_run_font(r, size=Pt(22), bold=True, color=RGBColor(0x1A,0x3C,0x6D))

sub = doc.add_paragraph(); sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = sub.add_run(f'数据范围：{date_range}    生成时间：{generated}    主题模式')
set_run_font(run, size=Pt(10), color=RGBColor(0x66,0x66,0x66))

total = sum(len(sec['items']) for sec in sections)
sub2 = doc.add_paragraph(); sub2.alignment = WD_ALIGN_PARAGRAPH.CENTER
run2 = sub2.add_run(f'共 {total} 条资讯 | {len(sections)} 大主题 | 覆盖教育部/人社部/人民日报/tech.net.cn/eol.cn/世校赛/各院校官网 | 严格06.10-06.16')
set_run_font(run2, size=Pt(9), color=RGBColor(0x99,0x99,0x99))
doc.add_paragraph('')

def add_hyperlink(paragraph, url, text):
    part = paragraph.part; r_id = part.relate_to(url, RT.HYPERLINK, is_external=True)
    hyperlink = OxmlElement('w:hyperlink'); hyperlink.set(qn('r:id'), r_id)
    new_run = OxmlElement('w:r'); rPr = OxmlElement('w:rPr')
    for tag, val in [('w:color','0563C1'),('w:u','single'),('w:sz','16')]:
        el = OxmlElement(tag); el.set(qn('w:val'), val); rPr.append(el)
    new_run.append(rPr); new_run.text = text; hyperlink.append(new_run)
    paragraph._p.append(hyperlink)

def hex_to_rgb(h):
    return RGBColor(int(h[0:2],16), int(h[2:4],16), int(h[4:6],16))

for sec in sections:
    h = doc.add_heading(sec['theme'], level=2)
    for r in h.runs: set_run_font(r, size=Pt(15), bold=True, color=hex_to_rgb(sec['color']))
    for item in sec['items']:
        p = doc.add_paragraph(); rt = p.add_run(f"▶ {item['title']}"); set_run_font(rt, size=Pt(11), bold=True)
        p2 = doc.add_paragraph()
        rd = p2.add_run(f"📅 {item['date']}    "); set_run_font(rd, size=Pt(9), color=RGBColor(0xE3,0x74,0x00))
        rs = p2.add_run(f"来源：{item['source']}"); set_run_font(rs, size=Pt(9), color=hex_to_rgb(sec['color']))
        pu = doc.add_paragraph(); add_hyperlink(pu, item['url'], '🔗 查看原文'); pu.paragraph_format.space_before = Pt(0)
        ps = doc.add_paragraph(); rsm = ps.add_run(item['summary']); set_run_font(rsm, size=Pt(11))
        ps.paragraph_format.left_indent = Cm(0.5)
        doc.add_paragraph('')

doc.add_paragraph('')
ps = doc.add_paragraph(); ps.alignment = WD_ALIGN_PARAGRAPH.CENTER
rs = ps.add_run(f'📊 共 {len(sections)} 大主题、{total} 条资讯 | 严格06.10-06.16 | 信息源：教育部/人社部/人民日报/中国教育报/光明日报/中国青年报/中国教育在线/人社部/各院校官网')
set_run_font(rs, size=Pt(9), color=RGBColor(0x66,0x66,0x66))

pn = doc.add_paragraph(); pn.alignment = WD_ALIGN_PARAGRAPH.CENTER
rn = pn.add_run('注：civte.edu.cn已curl验证可访问但本周无更新；cvae.com.cn为Vue SPA无法爬取。链接均可点击。')
set_run_font(rn, size=Pt(8), color=RGBColor(0x88,0x88,0x88))

fp = doc.add_paragraph(); fp.alignment = WD_ALIGN_PARAGRAPH.CENTER
rf = fp.add_run('🤖 主题模式 v3.0 | 基于《职教AI知识库-信息源网址与栏目映射》| Generated by Claude Code')
set_run_font(rf, size=Pt(8), color=RGBColor(0xAA,0xAA,0xAA))

out = "examples/education_news_thematic.docx"
doc.save(out)
print(f"✅ {out}")
print(f"{len(sections)} 主题、{total} 条（严格06.10-06.16）")
